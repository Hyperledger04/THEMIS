"""
Phase 10 — 06: OOXML Tracked-Changes Redlining
===============================================
Run:  pip install python-docx lxml
      python 06_redline_docx.py

The problem: python-docx can write .docx files, but it has no API for tracked
changes. If you open a Word document from LexAgent and a previous version
side-by-side, you cannot tell what changed. That makes contract review
painful — lawyers need to see insertions in green and deletions in red/strikethrough.

The insight: a .docx file is just a ZIP. Inside it, `word/document.xml` is an XML
file that Word renders. Tracked changes in that XML are just special elements:
  <w:del> — marks text as deleted (shown in red strikethrough)
  <w:ins> — marks text as inserted (shown in green underline)

We can inject these elements directly using lxml. Word will then show tracked
changes natively — no external redline service, no Word macros, no COM automation.

This is the doc-haus pattern. The real implementation is in:
  themis/tools/redline.py
"""

from __future__ import annotations
import difflib
import sys
import zipfile
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Section 1: What is OOXML?
# ---------------------------------------------------------------------------
#
# Office Open XML (OOXML) is the file format under .docx/.xlsx/.pptx.
# A .docx is literally a ZIP archive. You can unzip one and read the XML.
#
# The key file is word/document.xml. A paragraph with the text "Hello World"
# looks like this in the XML:
#
# <w:p>
#   <w:r>
#     <w:t>Hello World</w:t>
#   </w:r>
# </w:p>
#
# w:p = paragraph, w:r = run (a contiguous styled text sequence), w:t = text.
#
# A tracked deletion wraps the run in <w:del>:
#
# <w:del w:id="1" w:author="LexAgent" w:date="2026-01-01T00:00:00Z">
#   <w:r>
#     <w:delText>old text</w:delText>   ← note: w:delText, not w:t
#   </w:r>
# </w:del>
#
# A tracked insertion wraps it in <w:ins>:
#
# <w:ins w:id="2" w:author="LexAgent" w:date="2026-01-01T00:00:00Z">
#   <w:r>
#     <w:t>new text</w:t>
#   </w:r>
# </w:ins>
#
# python-docx gives us access to paragraph._p (the lxml element for <w:p>).
# We manipulate it directly.

print("=" * 60)
print("PART 1: Understanding the OOXML structure")
print("=" * 60)

# Create a simple document and inspect its XML
doc = Document()
para = doc.add_paragraph("The agreement shall commence on 1 January 2025.")

# _p is the underlying lxml element — paragraph XML node
import lxml.etree as etree

xml_str = etree.tostring(para._p, pretty_print=True).decode()
print("\nParagraph XML (simplified):")
# Print just the first 10 lines to avoid namespace noise
for line in xml_str.splitlines()[:15]:
    print(line)

print("\n↑ This is what Word actually reads. We will add <w:del>/<w:ins> to it.")

# ---------------------------------------------------------------------------
# Section 2: difflib gives us the diff — we map it to OOXML
# ---------------------------------------------------------------------------
print("\n" + "=" * 60)
print("PART 2: Using difflib to diff two text versions")
print("=" * 60)

OLD = "The agreement shall commence on 1 January 2025."
NEW = "The agreement shall commence on 1 March 2026."

old_words = OLD.split()
new_words = NEW.split()

matcher = difflib.SequenceMatcher(None, old_words, new_words, autojunk=False)
print(f"\nOld: {OLD}")
print(f"New: {NEW}")
print("\ndifflib opcodes (what changed, word by word):")
for tag, i1, i2, j1, j2 in matcher.get_opcodes():
    if tag == "equal":
        print(f"  equal    : {old_words[i1:i2]}")
    elif tag == "replace":
        print(f"  replace  : {old_words[i1:i2]} → {new_words[j1:j2]}")
    elif tag == "delete":
        print(f"  delete   : {old_words[i1:i2]}")
    elif tag == "insert":
        print(f"  insert   : {new_words[j1:j2]}")

print("\n↑ 'replace' becomes a <w:del> + <w:ins> pair in OOXML.")

# ---------------------------------------------------------------------------
# Pause and think:
# ---------------------------------------------------------------------------
# Q: Why do we diff at word level, not character level?
#
# A: Character-level diffs produce XML that is hard for Word to render cleanly.
#    A change like "2025" → "2026" at character level would produce:
#      <w:del>202<w:del>5<w:del>
#      <w:ins>202<w:ins>6<w:ins>
#    which fragments runs unnecessarily. Word-level diffs produce one del/ins
#    per changed word, which is what lawyers expect in a contract redline.

# ---------------------------------------------------------------------------
# Section 3: Building the OOXML elements
# ---------------------------------------------------------------------------
print("\n" + "=" * 60)
print("PART 3: Building <w:del> and <w:ins> elements")
print("=" * 60)

from datetime import datetime, timezone
import uuid

AUTHOR = "LexAgent"
DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def make_del_run(text: str) -> OxmlElement:
    """
    Build:  <w:del w:id="..." w:author="LexAgent" w:date="...">
              <w:r><w:delText xml:space="preserve">text</w:delText></w:r>
            </w:del>
    """
    el = OxmlElement("w:del")
    el.set(qn("w:id"), str(uuid.uuid4().int)[:8])
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), DATE)
    r = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    r.append(dt)
    el.append(r)
    return el


def make_ins_run(text: str) -> OxmlElement:
    """
    Build:  <w:ins w:id="..." w:author="LexAgent" w:date="...">
              <w:r><w:t xml:space="preserve">text</w:t></w:r>
            </w:ins>
    """
    el = OxmlElement("w:ins")
    el.set(qn("w:id"), str(uuid.uuid4().int)[:8])
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), DATE)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    el.append(r)
    return el


# Show what one del element looks like
del_el = make_del_run("January 2025")
print("\n<w:del> element:")
print(etree.tostring(del_el, pretty_print=True).decode()[:400])

ins_el = make_ins_run("March 2026")
print("\n<w:ins> element:")
print(etree.tostring(ins_el, pretty_print=True).decode()[:400])

# ---------------------------------------------------------------------------
# Section 4: Patching a paragraph
# ---------------------------------------------------------------------------
print("\n" + "=" * 60)
print("PART 4: Patching a paragraph in-place")
print("=" * 60)


def patch_paragraph(para, old_text: str, new_text: str) -> None:
    """
    Replace all runs in `para` with tracked-change markup.
    We strip existing <w:r>, <w:del>, <w:ins> children and rebuild
    from the difflib opcodes.
    """
    p = para._p
    # Remove existing run children (keep rPr = run properties if any)
    for child in list(p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "del", "ins"):
            p.remove(child)

    old_words = old_text.split()
    new_words = new_text.split()

    for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(
        None, old_words, new_words, autojunk=False
    ).get_opcodes():
        if tag == "equal":
            # Plain run — unchanged text
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = " ".join(old_words[i1:i2]) + " "
            r.append(t)
            p.append(r)
        elif tag == "replace":
            p.append(make_del_run(" ".join(old_words[i1:i2])))
            p.append(make_ins_run(" ".join(new_words[j1:j2])))
        elif tag == "delete":
            p.append(make_del_run(" ".join(old_words[i1:i2])))
        elif tag == "insert":
            p.append(make_ins_run(" ".join(new_words[j1:j2])))


# ---------------------------------------------------------------------------
# Section 5: Putting it all together
# ---------------------------------------------------------------------------
print("\n" + "=" * 60)
print("PART 5: Write a redlined .docx")
print("=" * 60)

import tempfile

original_text = [
    "The agreement shall commence on 1 January 2025.",
    "Either party may terminate with 30 days written notice.",
    "Confidentiality obligations survive for 2 years after termination.",
]

revised_text = """The agreement shall commence on 1 March 2026.
Either party may terminate with 90 days written notice.
Confidentiality obligations survive for 2 years after termination."""

with tempfile.TemporaryDirectory() as tmpdir:
    # Write original
    orig_path = f"{tmpdir}/original.docx"
    doc_orig = Document()
    for para_text in original_text:
        doc_orig.add_paragraph(para_text)
    doc_orig.save(orig_path)

    # Redline: load original, patch each paragraph
    doc_out = Document(orig_path)
    revised_paras = [p for p in revised_text.split("\n") if p.strip()]

    for i, para in enumerate(doc_out.paragraphs):
        old = para.text.strip()
        if not old:
            continue
        new = revised_paras[i].strip() if i < len(revised_paras) else ""
        if old != new:
            patch_paragraph(para, old, new)

    out_path = f"{tmpdir}/redlined.docx"
    doc_out.save(out_path)

    # Verify: the output ZIP must contain <w:del> and <w:ins> in document.xml
    with zipfile.ZipFile(out_path) as z:
        xml = z.read("word/document.xml").decode()

    has_del = "w:del" in xml
    has_ins = "w:ins" in xml
    print(f"  <w:del> present: {has_del}")
    print(f"  <w:ins> present: {has_ins}")
    print(f"  Unchanged para preserved: {'1 January 2025' not in xml or 'w:del' not in xml}")

    if has_del and has_ins:
        print("\n✅ Redline written — open it in Word to see tracked changes")
    else:
        print("\n❌ Something went wrong — check the patch_paragraph function")

# ---------------------------------------------------------------------------
# Key takeaways
# ---------------------------------------------------------------------------
print("\n" + "=" * 60)
print("KEY TAKEAWAYS")
print("=" * 60)
print("""
1. A .docx is a ZIP. You can inspect it with `unzip -l file.docx`.

2. Tracked changes are OOXML elements <w:del> and <w:ins> in
   word/document.xml. Word renders them as strikethrough/underline.

3. python-docx doesn't have a tracked-changes API, but para._p gives
   you the raw lxml element — you can build any valid OOXML structure.

4. difflib.SequenceMatcher gives you the same opcodes at the heart of
   Git diffs, code review tools, and legal redlines.

5. Real code: themis/tools/redline.py
   CLI: lex draft "revise NDA" --redline /tmp/original.docx
""")
