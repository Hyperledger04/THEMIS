"""
Phase 8 — UX & Output: 04_docx_writer.py
==========================================
Write court-ready .docx files from LexAgent's markdown-style draft text.

Install: pip install python-docx
Run:     python 04_docx_writer.py
         → creates /tmp/demo_writ.docx
"""

import os

from rich.console import Console
from rich.panel import Panel

console = Console()


# ── SECTION 1: WHY DOCX? ─────────────────────────────────────────────────────

# WRONG — save the draft as a plain .txt file.
# Indian courts and law offices expect Word documents (.docx).
# Plain text loses heading levels, spacing, and fonts — the clerk
# at the High Court filing counter will bounce it.
def wrong_save_as_text(draft: str, path: str) -> None:
    with open(path, "w") as f:
        f.write(draft)   # no formatting, no court-ready structure

# RIGHT — use python-docx to produce a properly structured Word document.
# LexAgent's write_docx() in lexagent/tools/docx_writer.py goes further:
# it sets Times New Roman 12pt, 1.5" left margin (High Court standard),
# double-spacing, justified alignment, and a citations appendix.
# Here we build the core write_docx() pattern step by step.

from docx import Document  # type: ignore[import]
from docx.shared import Pt  # type: ignore[import]


# ── SECTION 2: DOCUMENT BASICS ────────────────────────────────────────────────

def demo_document_basics():
    console.rule("[bold]SECTION 2 — Document Basics[/bold]")

    doc = Document()

    # doc.add_heading(text, level=N)
    #   level=0 → Title (largest)
    #   level=1 → Heading 1
    #   level=2 → Heading 2
    #   level=3 → Heading 3
    doc.add_heading("WRIT PETITION (CIVIL)", level=1)
    doc.add_heading("Ravi Kumar  v.  Union of India", level=2)

    # doc.add_paragraph(text) — plain body paragraph
    # Without a style, this produces a normal Word "Normal" paragraph.
    doc.add_paragraph("IN THE HON'BLE HIGH COURT OF DELHI AT NEW DELHI")
    doc.add_paragraph()  # blank spacer paragraph

    # doc.add_paragraph(text, style="List Bullet") — bulleted list
    doc.add_paragraph("Petitioner is a citizen of India.", style="List Bullet")
    doc.add_paragraph("Petitioner has been aggrieved by the impugned order.", style="List Bullet")
    doc.add_paragraph("This petition is maintainable under Article 226.", style="List Bullet")

    # doc.save(path) — write to disk
    path = "/tmp/demo_basics.docx"
    doc.save(path)
    console.print(f"[green]✓[/green] Saved basics demo: [underline]{path}[/underline]")
    console.print()


# ── SECTION 3: MARKDOWN-STYLE PARSER ─────────────────────────────────────────

# LexAgent's LLM returns draft text with lightweight markdown-style markup:
#   #  → Heading 1 (matter type)
#   ## → Heading 2 (sub-section: parties, prayer, grounds)
#   -  → Bullet point
#   (other) → Body paragraph
#
# write_docx() parses this format and maps it to Word document elements.

def write_docx(draft_text: str, output_path: str) -> None:
    """
    Convert markdown-style draft text into a court-ready .docx file.

    Parsing rules:
      Lines starting with '#  ' → Heading 1
      Lines starting with '## ' → Heading 2
      Lines starting with '-  ' → List Bullet paragraph
      Empty lines               → skipped (Word handles spacing via paragraph format)
      Everything else           → body paragraph (Normal style)

    This mirrors the actual LexAgent docx_writer at lexagent/tools/docx_writer.py,
    simplified to focus on the core parsing logic. The real version also handles:
      - Page margins (1.5" left — High Court standard)
      - Times New Roman 12pt font on all runs
      - Double-spacing (Pt(24) line spacing)
      - Justified alignment
      - Citations appendix
      - Matter metadata footer
    """
    doc = Document()

    for line in draft_text.splitlines():
        line = line.rstrip()   # remove trailing whitespace

        if not line:
            # Skip blank lines — python-docx paragraph spacing handles gaps
            continue

        if line.startswith("# "):
            # Heading level 1 — matter type, petition title
            heading_text = line[2:].strip()
            doc.add_heading(heading_text, level=1)

        elif line.startswith("## "):
            # Heading level 2 — sub-sections (PARTIES, GROUNDS, PRAYER)
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)

        elif line.startswith("- "):
            # Bulleted list item
            bullet_text = line[2:].strip()
            doc.add_paragraph(bullet_text, style="List Bullet")

        else:
            # Body paragraph — standard court pleading text
            doc.add_paragraph(line)

    abs_path = os.path.abspath(output_path)
    doc.save(abs_path)
    console.print(f"[green]✓[/green] Saved: [underline]{abs_path}[/underline]")


# ── SECTION 4: LIVE DEMO — WRIT PETITION SKELETON ────────────────────────────

# A 10-line writ petition skeleton that LexAgent would produce.
# Note the markdown-style headings and bullets — this is the format
# the LLM is prompted to return in lexagent/prompts/draft.md.
SAMPLE_WRIT_PETITION = """
# WRIT PETITION (CIVIL)

## Parties

Ravi Kumar, S/o Ram Kumar, R/o 12, Lajpat Nagar, New Delhi ... Petitioner
versus
Union of India, Through Secretary, Ministry of Home Affairs ... Respondent No. 1

## Grounds

- The impugned order dated 15.03.2024 is arbitrary, illegal, and violative of Article 14.
- The Petitioner was not afforded any opportunity of hearing before passing the order.
- The order violates the principles of natural justice as settled in Maneka Gandhi v. UOI.
- The Respondent has acted in excess of its statutory jurisdiction under Section 7 of the Act.

## Prayer

The Petitioner most humbly prays that this Hon'ble Court may be pleased to:
- Issue a Writ of Certiorari quashing the impugned order dated 15.03.2024.
- Issue a Writ of Mandamus directing the Respondent to reconsider the matter.
- Pass such other and further orders as this Hon'ble Court may deem fit and proper.
"""


def demo_writ_petition():
    console.rule("[bold]SECTION 4 — Live Demo: Writ Petition[/bold]")
    console.print("[dim]Creating /tmp/demo_writ.docx from markdown-style draft...[/dim]")
    console.print()

    write_docx(SAMPLE_WRIT_PETITION, "/tmp/demo_writ.docx")

    console.print()
    console.print(
        Panel(
            "[bold]What was parsed:[/bold]\n\n"
            "  [cyan]# WRIT PETITION (CIVIL)[/cyan]  → Heading 1\n"
            "  [cyan]## Parties[/cyan]               → Heading 2\n"
            "  [cyan]## Grounds[/cyan]               → Heading 2\n"
            "  [cyan]## Prayer[/cyan]                → Heading 2\n"
            "  [cyan]- The impugned order...[/cyan]  → List Bullet\n"
            "  [cyan]Ravi Kumar, S/o...[/cyan]       → Body paragraph\n\n"
            "[dim]Open /tmp/demo_writ.docx in Word or LibreOffice to verify.[/dim]",
            title="Parser Output",
            border_style="blue",
        )
    )
    console.print()


# ── SECTION 5: COMPARISON — WRONG VS RIGHT OUTPUT ────────────────────────────

def demo_wrong_vs_right():
    console.rule("[bold]SECTION 5 — Wrong vs Right Output[/bold]")

    # WRONG — plain text output
    wrong_path = "/tmp/demo_wrong.txt"
    with open(wrong_path, "w") as f:
        f.write(SAMPLE_WRIT_PETITION.strip())

    console.print(
        Panel(
            "[red]WRONG — Plain .txt file[/red]\n\n"
            "  - No heading hierarchy (# treated as literal character)\n"
            "  - No font settings (Times New Roman required by courts)\n"
            "  - No margins (1.5\" left margin is High Court standard)\n"
            "  - No justified alignment\n"
            "  - Clerk at filing counter rejects non-Word documents\n\n"
            f"  [dim]Saved to: {wrong_path}[/dim]",
            border_style="red",
        )
    )

    console.print()

    # RIGHT — structured .docx
    right_path = "/tmp/demo_right.docx"
    write_docx(SAMPLE_WRIT_PETITION, right_path)

    console.print(
        Panel(
            "[green]RIGHT — Structured .docx file[/green]\n\n"
            "  [green]✓[/green] Heading hierarchy preserved (H1 → H2)\n"
            "  [green]✓[/green] Bullet points formatted as Word list style\n"
            "  [green]✓[/green] Body paragraphs cleanly separated\n"
            "  [green]✓[/green] Ready for Times New Roman + margin formatting\n"
            "  [green]✓[/green] Word-compatible — opens in any court office\n\n"
            f"  [dim]Saved to: {right_path}[/dim]",
            border_style="green",
        )
    )
    console.print()


# ── MAIN ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    console.print()
    console.print(
        Panel(
            "[bold]Phase 8 — UX & Output[/bold]\n"
            "[dim]python-docx: writing court-ready Word documents[/dim]",
            border_style="blue",
        )
    )
    console.print()

    demo_document_basics()
    demo_writ_petition()
    demo_wrong_vs_right()

    console.print(
        "[dim]Next: open /tmp/demo_writ.docx to inspect the output. "
        "Then read lexagent/tools/docx_writer.py to see the full production version "
        "with margins, fonts, spacing, and citations appendix.[/dim]"
    )
    console.print()


# ── PAUSE AND THINK ──────────────────────────────────────────────────────────
#
# 1. Open lexagent/tools/docx_writer.py. The real write_docx() receives a
#    LexState dict, not a raw string. Why? What other state fields does it
#    use beyond draft_output? (Look for parties, jurisdiction, matter_id.)
#
# 2. lexagent/tools/docx_writer.py uses Inches(1.5) for the left margin.
#    Look up why Indian High Courts require a wider left margin — what
#    does the physical filing process do to that margin? How would this
#    break if the document were printed without margin settings?
#
# 3. The real write_docx() adds a Citations appendix (lines 86-94).
#    It reads from state["grounded_citations"] — a list of dicts with
#    "source", "verified", and "chunk_id" keys. Trace back to which
#    LexAgent node populates grounded_citations. (Hint: the cite node.)
#
# 4. python-docx doesn't support true footnotes without low-level XML
#    manipulation. The LexAgent team chose a citations appendix instead.
#    What is the tradeoff? Would a judge accept an appendix in place of
#    inline footnotes in a real Indian court filing?
#
# 5. The write_docx() parser here skips empty lines. The real version at
#    lexagent/tools/docx_writer.py uses re.split(r"\n\s*\n", draft) instead.
#    What's the difference? Which approach handles paragraphs separated by
#    two blank lines better?
