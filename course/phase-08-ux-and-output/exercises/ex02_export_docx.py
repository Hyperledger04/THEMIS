"""
Phase 8 — Exercise 02: Export Draft to .docx
=============================================
Implement write_docx() — a function that converts LexAgent's
markdown-style draft text into a court-ready Word document.

Install: pip install python-docx rich
Run:     python ex02_export_docx.py
         → creates /tmp/exercise_output.docx
"""

import os

from rich.console import Console
from rich.panel import Panel

console = Console()


# ── SAMPLE DRAFT ─────────────────────────────────────────────────────────────
# This is a 15-line writ petition draft in the markdown-style format
# that LexAgent's LLM returns. Your write_docx() must parse this correctly.

SAMPLE_DRAFT = """
# WRIT PETITION (CIVIL)

## Parties

Suresh Mehta, S/o Late Ramesh Mehta, R/o 45, Defence Colony, New Delhi ... Petitioner
versus
State of Delhi, Through Principal Secretary (Home) ... Respondent No. 1
Commissioner of Police, Delhi ... Respondent No. 2

## Grounds

- The impugned order dated 22.01.2024 is arbitrary, illegal, and without jurisdiction.
- The Petitioner was denied any opportunity of hearing in violation of natural justice.
- The action contravenes Article 14 of the Constitution as held in E.P. Royappa v. State of Tamil Nadu.
- The Respondents acted in excess of their authority under Section 144 CrPC.
- No reasoned order has been passed as mandated by the Supreme Court in S.N. Mukherjee v. UOI.

## Prayer

The Petitioner most humbly prays that this Hon'ble Court may be pleased to:
- Issue a Writ of Certiorari quashing the impugned order dated 22.01.2024.
- Issue a Writ of Mandamus directing the Respondents to restore the Petitioner's rights.
- Grant ad-interim stay of the impugned order pending disposal of this petition.
- Pass such other orders as this Hon'ble Court deems fit in the interests of justice.
"""


# ── YOUR TASK ─────────────────────────────────────────────────────────────────

def write_docx(draft_text: str, output_path: str) -> None:
    """
    Convert markdown-style draft text into a court-ready .docx file.

    Parsing rules — handle these four line types:
      Lines starting with '#  '  → doc.add_heading(text, level=1)
      Lines starting with '## '  → doc.add_heading(text, level=2)
      Lines starting with '-  '  → doc.add_paragraph(text, style="List Bullet")
      Empty lines                → skip (don't add blank paragraphs)
      Everything else            → doc.add_paragraph(text)   [body paragraph]

    After parsing, save the document to output_path and call:
        console.print(f"[green]✓[/green] Saved: [underline]{os.path.abspath(output_path)}[/underline]")

    Reference implementations:
      - lexagent/tools/docx_writer.py — the full production version
      - 04_docx_writer.py SECTION 3  — write_docx() teaching walkthrough

    Args:
        draft_text:  Raw draft string with # / ## / - markdown markers.
        output_path: File path for the .docx output (relative or absolute).
    """
    from docx import Document  # type: ignore[import]

    # TODO 1: Create a Document() and iterate over draft_text.splitlines().
    #         Strip trailing whitespace from each line.
    #         Skip lines that are empty after stripping.
    #         Hint: line.rstrip() cleans trailing spaces/newlines.
    pass  # replace with your implementation

    # TODO 2: For each non-empty line, check its prefix and call the
    #         correct Document method:
    #           line.startswith("# ")  → level=1 heading
    #           line.startswith("## ") → level=2 heading
    #           line.startswith("- ")  → List Bullet paragraph
    #           else                   → normal paragraph
    #         Remember to strip the prefix characters when passing text
    #         to the Document method (e.g., line[2:].strip() for "# " lines).
    pass  # replace with your implementation

    # TODO 3: Save the document with doc.save(output_path) and print the
    #         confirmation message shown in the docstring above.
    #         Use os.path.abspath(output_path) to show the full path.
    pass  # replace with your implementation


# ── DRIVER ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    console.print()
    console.print(
        Panel(
            "[bold]Exercise 02 — Export Draft to .docx[/bold]\n"
            "[dim]Implement write_docx() with markdown-to-Word parsing[/dim]",
            border_style="blue",
        )
    )
    console.print()

    # ── Test 1: Write the sample writ petition ────────────────────────────────
    console.rule("[bold]Test 1 — Sample Writ Petition[/bold]")
    write_docx(SAMPLE_DRAFT, "/tmp/exercise_output.docx")

    console.print()
    console.print(
        "[dim]Open /tmp/exercise_output.docx to verify:\n"
        "  - 'WRIT PETITION (CIVIL)' should be a large Heading 1\n"
        "  - 'Parties', 'Grounds', 'Prayer' should be Heading 2\n"
        "  - Bullet points should be List Bullet style (with bullet symbol)\n"
        "  - Body text should be plain paragraphs[/dim]"
    )

    console.print()

    # ── Test 2: Edge case — only body paragraphs, no headings ────────────────
    console.rule("[bold]Test 2 — Body Only (no headings)[/bold]")
    body_only = """
The Petitioner respectfully submits this petition under Article 226.
The facts of the case are as stated in the accompanying affidavit.
Relief is sought as prayed herein below.
"""
    write_docx(body_only, "/tmp/exercise_body_only.docx")

    console.print()

    # ── Test 3: Bullets only ──────────────────────────────────────────────────
    console.rule("[bold]Test 3 — Bullets Only[/bold]")
    bullets_only = """
- Relief 1: Quash the impugned order.
- Relief 2: Issue Mandamus to Respondent.
- Relief 3: Award costs of this petition.
"""
    write_docx(bullets_only, "/tmp/exercise_bullets.docx")

    console.print()

    # ── Test 4: Empty input — should not crash ────────────────────────────────
    console.rule("[bold]Test 4 — Empty Input (no crash)[/bold]")
    write_docx("", "/tmp/exercise_empty.docx")

    console.print()
    console.print(
        Panel(
            "[bold green]All tests done.[/bold green]\n\n"
            "  Open the .docx files in Word or LibreOffice to inspect the output.\n"
            "  Then read [underline]lexagent/tools/docx_writer.py[/underline] to see\n"
            "  how the production version adds margins, font, spacing, and citations.",
            border_style="green",
        )
    )
    console.print()


# ── REFLECTION QUESTIONS ──────────────────────────────────────────────────────
#
# Q1. Open lexagent/tools/docx_writer.py. The production write_docx() uses
#     re.split(r"\n\s*\n", draft) instead of splitlines() to split into
#     paragraphs. What's the difference? Which approach handles a paragraph
#     that spans multiple lines without a heading marker?
#     (Hint: try a draft with a 3-line paragraph and no blank lines between them.)
#
# Q2. The production docx_writer.py calls _set_font(para) and
#     _set_paragraph_spacing(para) after creating each paragraph.
#     Why must these be called AFTER doc.add_paragraph() rather than before?
#     (Hint: look at what _set_font() iterates over — para.runs.)
#
# Q3. Your write_docx() uses os.path.abspath(output_path) before saving.
#     The real LexAgent version returns this absolute path as a string.
#     Why is returning the absolute path useful to the caller — specifically,
#     when LexAgent's Telegram bot sends the file back to the lawyer?
