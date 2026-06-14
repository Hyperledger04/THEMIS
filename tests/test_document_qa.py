# Tests for the document QA node — PDF/DOCX parsing, chunking, and retrieval.
# Uses in-memory data only; no real PDF/DOCX files needed for unit tests.

import pytest
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

from themis.nodes.document_qa import (
    DocChunk,
    _find_doc_chunk,
    _infer_location,
    build_retriever,
    parse_docx,
)


# ---------------------------------------------------------------------------
# _infer_location
# ---------------------------------------------------------------------------

def test_infer_location_clause():
    text = "Clause 12.3 The licensee shall not sub-license any rights."
    loc = _infer_location(text, page=4)
    assert "Page 4" in loc
    assert "Clause 12.3" in loc


def test_infer_location_section():
    text = "Section 14B — Penalties and remedies applicable hereunder."
    loc = _infer_location(text, page=7)
    assert "Section 14" in loc
    assert "Page 7" in loc


def test_infer_location_no_label():
    text = "This agreement is entered into between the parties."
    loc = _infer_location(text, page=1)
    assert loc == "Page 1"


# ---------------------------------------------------------------------------
# _find_doc_chunk
# ---------------------------------------------------------------------------

def test_find_doc_chunk_exact_match():
    chunks = [
        DocChunk(text="The licensee shall not sub-license.", page=4, location="Page 4, Clause 12", chunk_index=0),
        DocChunk(text="Governing law: Delhi courts.", page=8, location="Page 8, Clause 18", chunk_index=1),
    ]
    found = _find_doc_chunk("The licensee shall not sub-license.", chunks)
    assert found is not None
    assert found.chunk_index == 0


def test_find_doc_chunk_partial_match():
    chunks = [
        DocChunk(text="The licensee shall not sub-license any rights without consent.", page=4, location="Page 4", chunk_index=0),
    ]
    found = _find_doc_chunk("licensee shall not sub-license", chunks)
    assert found is not None


def test_find_doc_chunk_not_found():
    chunks = [
        DocChunk(text="Completely unrelated text.", page=1, location="Page 1", chunk_index=0),
    ]
    found = _find_doc_chunk("Something entirely different", chunks)
    assert found is None


# ---------------------------------------------------------------------------
# build_retriever
# ---------------------------------------------------------------------------

def test_build_retriever_returns_hybrid_retriever():
    from themis.tools.retriever import HybridRetriever
    chunks = [
        DocChunk(text="The agreement shall be governed by the laws of India.", page=1, location="Page 1, Clause 1", chunk_index=0),
        DocChunk(text="Either party may terminate on 30 days' written notice.", page=3, location="Page 3, Clause 8", chunk_index=1),
        DocChunk(text="Confidential information shall not be disclosed.", page=5, location="Page 5, Clause 15", chunk_index=2),
    ]
    retriever = build_retriever(chunks)
    assert isinstance(retriever, HybridRetriever)


def test_build_retriever_retrieves_relevant():
    # Use a larger corpus so BM25+TF-IDF scores are meaningful
    chunks = [
        DocChunk(text="The agreement shall be governed by the laws of India and jurisdiction of Delhi courts.", page=1, location="Page 1, Clause 1", chunk_index=0),
        DocChunk(text="Confidential information shall not be disclosed to third parties without prior written consent.", page=5, location="Page 5, Clause 15", chunk_index=1),
        DocChunk(text="Payment shall be due within 30 days of invoice receipt from the creditor.", page=2, location="Page 2, Clause 5", chunk_index=2),
        DocChunk(text="The licensee may not sub-license rights granted under this agreement.", page=3, location="Page 3, Clause 8", chunk_index=3),
        DocChunk(text="Termination shall be effected by written notice to the other party.", page=4, location="Page 4, Clause 12", chunk_index=4),
    ]
    retriever = build_retriever(chunks)
    results = retriever.retrieve("governing law India jurisdiction", top_k=3)
    # Retriever returns at least one result for this query
    assert len(results) >= 1


# ---------------------------------------------------------------------------
# parse_docx (mocked)
# ---------------------------------------------------------------------------

def test_parse_docx_basic(tmp_path):
    """Verify DOCX parsing extracts paragraphs into DocChunks."""
    from docx import Document

    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This agreement is entered into between Party A and Party B.")
    doc.add_heading("Clause 5 — Payment", level=2)
    doc.add_paragraph("Payment shall be made within 30 days of the invoice date.")
    doc.save(str(doc_path))

    chunks = parse_docx(doc_path)
    assert len(chunks) >= 1
    all_text = " ".join(c.text for c in chunks)
    assert "Party A" in all_text or "agreement" in all_text


def test_parse_docx_empty_paragraphs(tmp_path):
    """Empty paragraphs produce no chunks."""
    from docx import Document

    doc_path = tmp_path / "empty.docx"
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.save(str(doc_path))

    chunks = parse_docx(doc_path)
    assert len(chunks) == 0


# ---------------------------------------------------------------------------
# answer_question (mocked LLM)
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_answer_question_returns_inline_citations():
    from themis.nodes.document_qa import answer_question
    from themis.config import LexConfig

    chunks = [
        DocChunk(text="The licensee shall not sub-license any rights without prior written consent.", page=4, location="Page 4, Clause 12.3", chunk_index=0),
        DocChunk(text="This agreement is governed by the laws of India.", page=8, location="Page 8, Clause 18", chunk_index=1),
        DocChunk(text="Payment is due within 30 days of invoice receipt.", page=2, location="Page 2, Clause 5", chunk_index=2),
    ]
    retriever = build_retriever(chunks)
    cfg = LexConfig()

    with patch("themis.nodes._llm.call_llm", new_callable=AsyncMock) as mock_call_llm:
        mock_call_llm.return_value = {
            "content": "Sub-licensing is prohibited [1]. Governing law is India [2].",
            "tool_calls": None,
        }

        answer, cited = await answer_question("Can the licensee sub-license?", chunks, retriever, cfg, top_k=3)

    assert "[1]" in answer or "[2]" in answer or "Sub-licensing" in answer
    assert isinstance(cited, list)
    assert len(cited) > 0


@pytest.mark.asyncio
async def test_answer_question_no_relevant_chunks():
    from themis.nodes.document_qa import answer_question
    from themis.config import LexConfig

    # Empty document
    chunks: list[DocChunk] = []
    retriever = build_retriever(chunks)
    cfg = LexConfig()

    answer, cited = await answer_question("What is the governing law?", chunks, retriever, cfg)
    assert "does not appear" in answer.lower() or answer == ""
    assert cited == []
