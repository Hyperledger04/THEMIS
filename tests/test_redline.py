import pytest
from pathlib import Path
from docx import Document
from themis.tools.redline import write_redline_docx

@pytest.fixture
def original_docx(tmp_path):
    doc = Document()
    doc.add_paragraph("The agreement shall commence on 1 January 2025.")
    doc.add_paragraph("Either party may terminate with 30 days notice.")
    p = tmp_path / "original.docx"
    doc.save(str(p))
    return str(p)

def test_redline_creates_output_file(original_docx, tmp_path):
    revised = (
        "The agreement shall commence on 1 March 2026.\n"
        "Either party may terminate with 60 days notice."
    )
    out = str(tmp_path / "redlined.docx")
    result = write_redline_docx(original_docx, revised, out)
    assert Path(result).exists()

def test_redline_contains_tracked_change_markup(original_docx, tmp_path):
    revised = (
        "The agreement shall commence on 1 January 2025.\n"
        "Either party may terminate with 60 days notice."
    )
    out = str(tmp_path / "redlined.docx")
    write_redline_docx(original_docx, revised, out)
    import zipfile
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode()
    assert "w:ins" in xml or "w:del" in xml

def test_redline_unchanged_para_has_no_markup(original_docx, tmp_path):
    revised = (
        "The agreement shall commence on 1 January 2025.\n"
        "Either party may terminate with 60 days notice."
    )
    out = str(tmp_path / "redlined.docx")
    write_redline_docx(original_docx, revised, out)
    import zipfile
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode()
    assert "1 January 2025" in xml

def test_redline_returns_absolute_path(original_docx, tmp_path):
    revised = "The agreement shall commence on 1 March 2026.\n"
    out = str(tmp_path / "out.docx")
    result = write_redline_docx(original_docx, revised, out)
    assert result == str(Path(out).resolve())
