"""Tests for lexagent/tools/docx_writer.py"""
import os
import pytest
from docx import Document
from lexagent.tools.docx_writer import write_docx


def _state(**overrides) -> dict:
    base = {
        "draft_output": "This is paragraph one.\n\nThis is paragraph two with more content.",
        "matter_type": "injunction",
        "parties": {"plaintiff": "Ram Lal", "defendant": "Shyam Builders"},
        "jurisdiction": "Delhi High Court",
        "matter_id": "M-DOC-001",
        "grounded_citations": [
            {"source": "AIR 1978 SC 597", "chunk_id": "AIR 1978 SC 597::0", "verified": True},
            {"source": "(2021) 3 SCC 415", "chunk_id": None, "verified": False},
        ],
    }
    base.update(overrides)
    return base


def test_write_docx_creates_file(tmp_path):
    out = str(tmp_path / "out.docx")
    result = write_docx(_state(), out)
    assert os.path.exists(result)
    assert result.endswith(".docx")


def test_write_docx_returns_absolute_path(tmp_path):
    out = str(tmp_path / "draft.docx")
    result = write_docx(_state(), out)
    assert os.path.isabs(result)


def test_write_docx_readable_by_python_docx(tmp_path):
    out = str(tmp_path / "readable.docx")
    write_docx(_state(), out)
    doc = Document(out)
    # Document must have at least one paragraph
    assert len(doc.paragraphs) > 0


def test_write_docx_contains_parties(tmp_path):
    out = str(tmp_path / "parties.docx")
    write_docx(_state(), out)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Ram Lal" in full_text
    assert "Shyam Builders" in full_text


def test_write_docx_contains_draft_text(tmp_path):
    out = str(tmp_path / "draft_content.docx")
    write_docx(_state(), out)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "paragraph one" in full_text


def test_write_docx_contains_citations_appendix(tmp_path):
    out = str(tmp_path / "citations.docx")
    write_docx(_state(), out)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "AIR 1978 SC 597" in full_text


def test_write_docx_no_parties_uses_defaults(tmp_path):
    out = str(tmp_path / "noparties.docx")
    state = _state(parties=None)
    result = write_docx(state, out)
    assert os.path.exists(result)


def test_write_docx_no_jurisdiction(tmp_path):
    out = str(tmp_path / "nojur.docx")
    state = _state(jurisdiction=None)
    result = write_docx(state, out)
    assert os.path.exists(result)


def test_write_docx_no_grounded_citations(tmp_path):
    out = str(tmp_path / "nocites.docx")
    state = _state(grounded_citations=None)
    result = write_docx(state, out)
    assert os.path.exists(result)


def test_write_docx_matter_id_in_footer(tmp_path):
    out = str(tmp_path / "footer.docx")
    write_docx(_state(), out)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "M-DOC-001" in full_text
