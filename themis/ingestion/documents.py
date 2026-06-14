"""
Document ingestion pipeline.

Accepts a local file path (PDF, DOCX, TXT, plain text), extracts text
page-by-page, creates SourceAnchors for every non-blank line, persists a
DocumentRecord, and returns structured page text for downstream extraction.

MVP storage: local filesystem under ~/.themis/matters/{matter_id}/documents/.
Scale path: swap storage_uri to S3/R2 by changing _store_file only.

OCR fallback (for scanned images/PDFs with no embedded text) is deferred
to ingestion/ocr.py — this module calls it only when pdfplumber returns
empty text for a page.
"""
from __future__ import annotations

import hashlib
import logging
import mimetypes
import os
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from themis.ingestion.anchors import build_line_anchors
from themis.workspace.models import DocumentRecord

logger = logging.getLogger(__name__)

# Base storage directory for all matter documents (overridable via env)
_STORAGE_ROOT = Path(os.environ.get("THEMIS_STORAGE_ROOT", Path.home() / ".themis" / "documents"))


@dataclass
class PageText:
    """One page of extracted text with its page number (1-indexed)."""
    page: int
    text: str
    char_count: int


@dataclass
class IngestedDocument:
    """Result returned after full ingestion of a single file."""
    record: DocumentRecord
    pages: list[PageText]
    anchor_count: int


def ingest_file(
    file_path: str | Path,
    matter_id: str,
    firm_id: str,
    repo,
    run_id: Optional[str] = None,
) -> IngestedDocument:
    """
    Full ingestion pipeline for a single file:
      1. Copy file to local storage with a content-addressed name.
      2. Detect mime type.
      3. Extract text page-by-page.
      4. Build SourceAnchors from line text.
      5. Persist DocumentRecord + anchors via repo.
      6. Return IngestedDocument with page text for downstream extraction.

    Args:
        file_path: Path to the source file (can be a temp upload location).
        matter_id: Target matter.
        firm_id: Owning firm (tenant scope).
        repo: PostgresWorkspaceRepository (or compatible).
        run_id: Optional extraction run ID for provenance.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"Ingestion file not found: {file_path}")

    storage_uri, mime_type = _store_file(file_path, matter_id, firm_id)
    parser, pages = _extract_text(file_path, mime_type)

    doc = DocumentRecord(
        matter_id=matter_id,
        firm_id=firm_id,
        filename=file_path.name,
        mime_type=mime_type,
        storage_uri=storage_uri,
        parser=parser,
        page_count=len(pages),
        status="indexed",
    )
    repo.create_document(doc)

    # Build SourceAnchors for every non-blank line across all pages
    all_anchors = []
    for page in pages:
        anchors = build_line_anchors(
            matter_id=matter_id,
            document_id=doc.document_id,
            page=page.page,
            text=page.text,
            extraction_run_id=run_id,
        )
        all_anchors.extend(anchors)

    repo.bulk_create_anchors(all_anchors)
    logger.info(
        "Ingested %s: parser=%s pages=%d anchors=%d",
        file_path.name, parser, len(pages), len(all_anchors),
    )
    return IngestedDocument(record=doc, pages=pages, anchor_count=len(all_anchors))


# ------------------------------------------------------------------
# Internal helpers
# ------------------------------------------------------------------

def _store_file(file_path: Path, matter_id: str, firm_id: str) -> tuple[str, str]:
    """
    Copy file to content-addressed local storage.
    Returns (storage_uri, mime_type).
    """
    dest_dir = _STORAGE_ROOT / firm_id / matter_id
    dest_dir.mkdir(parents=True, exist_ok=True)

    # Content-addressed name prevents duplicates on re-upload
    file_hash = _sha256_prefix(file_path)
    dest_name = f"{file_hash}_{file_path.name}"
    dest_path = dest_dir / dest_name

    if not dest_path.exists():
        shutil.copy2(file_path, dest_path)

    storage_uri = str(dest_path)
    mime_type = mimetypes.guess_type(file_path.name)[0] or "application/octet-stream"
    return storage_uri, mime_type


def _extract_text(file_path: Path, mime_type: str) -> tuple[str, list[PageText]]:
    """
    Dispatch to the appropriate extractor based on mime type / extension.
    Returns (parser_name, list[PageText]).
    """
    suffix = file_path.suffix.lower()

    if mime_type == "application/pdf" or suffix == ".pdf":
        return _extract_pdf(file_path)

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or suffix in (".docx", ".doc"):
        return _extract_docx(file_path)

    if mime_type and mime_type.startswith("text/") or suffix in (".txt", ".md", ".rst"):
        return _extract_plaintext(file_path)

    # Unknown type — attempt plaintext as last resort
    logger.warning("Unknown mime type %s for %s — trying plaintext.", mime_type, file_path.name)
    return _extract_plaintext(file_path)


def _extract_pdf(file_path: Path) -> tuple[str, list[PageText]]:
    """Extract text from a PDF using pdfplumber, one PageText per page."""
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("pdfplumber is required for PDF ingestion. Run: uv add pdfplumber") from exc

    pages: list[PageText] = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                # Attempt OCR fallback for image-only pages
                text = _try_ocr_page(file_path, page_number=i)
            pages.append(PageText(page=i, text=text, char_count=len(text)))

    return "pdfplumber", pages


def _extract_docx(file_path: Path) -> tuple[str, list[PageText]]:
    """
    Extract text from a DOCX file.
    DOCX has no native page concept — we treat the whole document as page 1,
    then split on paragraph boundaries at ~3000-char chunks as synthetic pages.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX ingestion. Run: uv add python-docx") from exc

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    # Synthetic pagination at ~3000 chars so anchors stay within readable page windows
    chunk_size = 3000
    chunks = [full_text[i : i + chunk_size] for i in range(0, len(full_text), chunk_size)] or [""]
    pages = [
        PageText(page=i + 1, text=chunk, char_count=len(chunk))
        for i, chunk in enumerate(chunks)
    ]
    return "docx", pages


def _extract_plaintext(file_path: Path) -> tuple[str, list[PageText]]:
    """Read a plain-text file and split into synthetic pages at ~3000 chars."""
    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        logger.warning("Could not read %s as text: %s", file_path, exc)
        return "unknown", [PageText(page=1, text="", char_count=0)]

    chunk_size = 3000
    chunks = [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)] or [""]
    pages = [
        PageText(page=i + 1, text=chunk, char_count=len(chunk))
        for i, chunk in enumerate(chunks)
    ]
    return "plaintext", pages


def _try_ocr_page(file_path: Path, page_number: int) -> str:
    """
    Attempt OCR on a single PDF page that returned no embedded text.
    Returns empty string if pytesseract is not installed — OCR is optional.
    """
    try:
        from themis.ingestion.ocr import ocr_pdf_page
        return ocr_pdf_page(file_path, page_number)
    except ImportError:
        return ""


def _sha256_prefix(file_path: Path, prefix_len: int = 12) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()[:prefix_len]
