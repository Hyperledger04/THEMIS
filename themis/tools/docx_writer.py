# Court-ready .docx formatter using python-docx.
#
# Output contract (court-fileable):
#   - Formal court header: all-caps, bold, centered plain paragraph
#   - Cause title: party labels driven by CourtDraftSpec (falls back to PARTY_LABELS dict)
#   - Body: draft paragraphs (spacing/font from spec) — lawyer notes stripped
#   - Citations appendix: grounded citations with source chunk reference
#   - lawyer_notes.docx: Plain English Summary + Risk Assessment routed here, NOT in the filing
#
# Formatting is driven by CourtDraftSpec loaded from themis/courts/*.yaml.
# The spec is resolved from (jurisdiction, matter_type) + lawyer's SOUL.md formatting_style.

from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document  # type: ignore[import]
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # type: ignore[import]
from docx.shared import Inches, Mm, Pt  # type: ignore[import]

from themis.courts import CourtDraftSpec, load_spec
from themis.state import SeniorCounselState

# ---------------------------------------------------------------------------
# Party label map — keyed by matter_type value from LexState.
# WHY: "Petitioner v. Respondent" is writ-petition vocabulary.
# An S.138 criminal complaint must say "Complainant ... Accused" or the
# court officer will flag a format defect at the filing counter.
# Default falls back to Petitioner/Respondent for unknown matter types.
# ---------------------------------------------------------------------------
PARTY_LABELS: dict[str, tuple[str, str]] = {
    "s138_complaint":         ("Complainant", "Accused"),
    "criminal_complaint":     ("Complainant", "Accused"),
    "bail_application":       ("Accused", "State"),
    "writ_petition":          ("Petitioner", "Respondent"),
    "plaint":                 ("Plaintiff", "Defendant"),
    "written_statement":      ("Plaintiff", "Defendant"),
    "injunction_application": ("Plaintiff/Applicant", "Defendant/Respondent"),
    "legal_notice":           ("Sender", "Addressee"),
    "civil_suit":             ("Plaintiff", "Defendant"),
}
_DEFAULT_PARTY_LABELS = ("Petitioner", "Respondent")


def _party_labels(matter_type: str | None) -> tuple[str, str]:
    if not matter_type:
        return _DEFAULT_PARTY_LABELS
    key = (matter_type or "").lower().strip().replace(" ", "_")
    # Exact match first, then substring scan
    if key in PARTY_LABELS:
        return PARTY_LABELS[key]
    for k, v in PARTY_LABELS.items():
        if k in key or key in k:
            return v
    return _DEFAULT_PARTY_LABELS


def _filer_party(matter_type: str | None, parties: dict) -> str:
    """Return the filer-side party name from the parties dict using the correct key."""
    label_filer, _ = _party_labels(matter_type)
    label_key = label_filer.lower()
    # Try exact key match, then alias keys
    return (
        parties.get(label_key)
        or parties.get("plaintiff")
        or parties.get("complainant")
        or parties.get("petitioner")
        or parties.get("accused")
        or label_filer
    )


def _opposing_party(matter_type: str | None, parties: dict) -> str:
    """Return the opposing-side party name from the parties dict."""
    _, label_opposing = _party_labels(matter_type)
    label_key = label_opposing.lower()
    return (
        parties.get(label_key)
        or parties.get("defendant")
        or parties.get("accused")
        or parties.get("respondent")
        or label_opposing
    )


# ---------------------------------------------------------------------------
# Court header renderer
# ---------------------------------------------------------------------------

def _format_court_header(jurisdiction: str) -> str:
    """
    Convert a jurisdiction string to a formal court header.

    Examples:
      "GBN CJM" → "IN THE COURT OF THE CHIEF JUDICIAL MAGISTRATE, GAUTAM BUDH NAGAR"
      "Delhi High Court" → "IN THE HIGH COURT OF DELHI AT NEW DELHI"
      "IN THE COURT OF ..." → returned unchanged (already formal)

    If the string already starts with "IN THE", it is returned as-is, uppercased.
    Otherwise a best-effort expansion is applied; raw string is uppercased as fallback.
    """
    j = jurisdiction.strip()
    if j.upper().startswith("IN THE"):
        return j.upper()

    # Common shorthand expansions
    expansions = {
        "gbncjm":             "IN THE COURT OF THE CHIEF JUDICIAL MAGISTRATE, GAUTAM BUDH NAGAR AT GREATER NOIDA",
        "gbn cjm":            "IN THE COURT OF THE CHIEF JUDICIAL MAGISTRATE, GAUTAM BUDH NAGAR AT GREATER NOIDA",
        "delhi hc":           "IN THE HIGH COURT OF DELHI AT NEW DELHI",
        "delhi high court":   "IN THE HIGH COURT OF DELHI AT NEW DELHI",
        "bombay hc":          "IN THE HIGH COURT OF JUDICATURE AT BOMBAY",
        "bombay high court":  "IN THE HIGH COURT OF JUDICATURE AT BOMBAY",
        "madras hc":          "IN THE HIGH COURT OF JUDICATURE AT MADRAS",
        "allahabad hc":       "IN THE HIGH COURT OF JUDICATURE AT ALLAHABAD",
        "supreme court":      "IN THE SUPREME COURT OF INDIA",
    }
    key = j.lower().strip()
    if key in expansions:
        return expansions[key]

    # If it looks like "Chief Judicial Magistrate, XYZ" wrap it
    if "chief judicial magistrate" in key or "cjm" in key:
        return f"IN THE COURT OF THE {j.upper()}"
    if "high court" in key:
        return f"IN THE {j.upper()}"
    if "sessions court" in key or "sessions judge" in key:
        return f"IN THE COURT OF THE {j.upper()}"

    # Fallback: uppercase the raw string — still looks formal
    return j.upper()


def _add_court_header(
    doc: Document,
    court_header: str,
    font_name: str = "Bookman Old Style",
    font_size_pt: int = 14,
) -> None:
    """Write the court header as a bold, centered plain paragraph (not a heading style)."""
    for line in court_header.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)


# ---------------------------------------------------------------------------
# Draft splitter — separates court-filing body from lawyer working notes
# ---------------------------------------------------------------------------

def _split_draft(draft: str) -> tuple[str, str]:
    """
    Split the LLM draft at the first '---' separator line.

    Returns (filing_body, lawyer_notes).
    - filing_body:    everything before the separator → goes into the .docx filing
    - lawyer_notes:   everything after the separator  → goes into lawyer_notes.docx

    The separator pattern is a line that is exactly '---' (possibly with surrounding
    whitespace), which is the convention used in base_system.md to delimit the
    Plain English Summary and Risk Assessment from the legal document body.
    """
    # Match a line that is ONLY dashes (3+), possibly with surrounding whitespace
    parts = re.split(r"\n[ \t]*---+[ \t]*\n", draft, maxsplit=1)
    if len(parts) == 2:
        return parts[0].rstrip(), parts[1].strip()
    return draft, ""


# ---------------------------------------------------------------------------
# Lawyer notes writer — separate .docx for internal use only
# ---------------------------------------------------------------------------

def _write_lawyer_notes_docx(notes_text: str, notes_path: str) -> str:
    """Write lawyer notes (Plain English Summary + Risk flags) to a separate .docx."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("LAWYER'S WORKING NOTES — NOT FOR FILING")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    for para_text in [p.strip() for p in re.split(r"\n\s*\n", notes_text) if p.strip()]:
        p = doc.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_font(p)
        _set_paragraph_spacing(p)

    abs_path = os.path.abspath(notes_path)
    doc.save(abs_path)
    return abs_path


# ---------------------------------------------------------------------------
# Main public function
# ---------------------------------------------------------------------------

def write_docx(state: SeniorCounselState, output_path: str) -> str:
    """
    Write state["draft_output"] to a court-ready .docx file.

    - Loads CourtDraftSpec from (jurisdiction, matter_type, soul) to drive
      page size, margins, font, spacing, and party labels.
    - Splits draft at '---': filing body → output_path; lawyer notes → <stem>_notes.docx.
    - Renders a formal court header from jurisdiction string.

    Returns the resolved absolute path of the filing .docx.
    """
    matter_type = state.get("matter_type") or ""
    parties = state.get("parties") or {}
    jurisdiction = state.get("jurisdiction") or ""
    soul = state.get("lawyer_soul") or {}
    soul_dict = soul if isinstance(soul, dict) else {}

    # ----------------------------------------------------------------
    # Load court spec — drives all formatting decisions from here on.
    # WHY: all hardcoded font/margin/spacing values are gone; the spec
    # encodes the per-court rules so this function stays court-agnostic.
    # ----------------------------------------------------------------
    spec = load_spec(jurisdiction, matter_type, soul=soul_dict)
    fmt = spec.formatting

    doc = Document()
    _apply_page_setup(doc, spec)

    # ----------------------------------------------------------------
    # Court header — formal all-caps centered paragraph, not a heading
    # WHY: Try SOUL.md court_name_formal first; fall back to jurisdiction string.
    # ----------------------------------------------------------------
    court_header = soul_dict.get("court_name_formal") or ""
    if not court_header and jurisdiction:
        court_header = _format_court_header(jurisdiction)
    if court_header:
        _add_court_header(doc, court_header, font_name=fmt.font_name, font_size_pt=fmt.font_size_pt)
        doc.add_paragraph()

    # ----------------------------------------------------------------
    # Blank case number line — court fills this at filing counter
    # ----------------------------------------------------------------
    case_no_p = doc.add_paragraph()
    case_no_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = case_no_p.add_run("COMPLAINT CASE NO. _____ OF ____")
    run.bold = True
    run.font.name = fmt.font_name
    run.font.size = Pt(fmt.font_size_pt)
    doc.add_paragraph()

    # ----------------------------------------------------------------
    # Cause title — party labels from spec (court-specific), falling back
    # to the legacy PARTY_LABELS dict for matter types without a YAML.
    # WHY: spec.party_labels is ("Complainant", "Accused") for S.138 CJM
    # and ("Petitioner", "Respondent") for HC writs — from the YAML.
    # ----------------------------------------------------------------
    spec_filer_label, spec_opposing_label = spec.party_labels
    # prefer spec labels; _party_labels() is a fallback for unmatched types
    filer_label = spec_filer_label or _party_labels(matter_type)[0]
    opposing_label = spec_opposing_label or _party_labels(matter_type)[1]
    filer_name = _filer_party(matter_type, parties)
    opposing_name = _opposing_party(matter_type, parties)

    for text, bold in [
        (filer_name, True),
        (f"... {filer_label.upper()}", False),
        ("VERSUS", True),
        (opposing_name, True),
        (f"... {opposing_label.upper()}", False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = fmt.font_name
        run.font.size = Pt(fmt.font_size_pt)

    doc.add_paragraph()

    # ----------------------------------------------------------------
    # Split draft at '---' to separate filing body from lawyer notes
    # WHY: base_system.md instructs the LLM to append Plain English Summary
    # and Risk Assessment after a '---' separator. These must never appear
    # in a court filing.
    # ----------------------------------------------------------------
    raw_draft = state.get("draft_output") or ""
    filing_body, lawyer_notes_text = _split_draft(raw_draft)

    for para_text in [p.strip() for p in re.split(r"\n\s*\n", filing_body) if p.strip()]:
        p = doc.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_font(p, font_name=fmt.font_name, size=fmt.font_size_pt)
        _set_paragraph_spacing(p, line_spacing=fmt.line_spacing)

    # ----------------------------------------------------------------
    # Citations appendix — grounded citations only (no lawyer flags)
    # ----------------------------------------------------------------
    grounded = state.get("grounded_citations") or []
    if grounded:
        doc.add_page_break()
        heading_p = doc.add_paragraph()
        heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading_p.add_run("LIST OF CITATIONS")
        run.bold = True
        run.font.name = fmt.font_name
        run.font.size = Pt(fmt.font_size_pt)

        for i, g in enumerate(grounded, start=1):
            status = "✓ Verified" if g["verified"] else "⚠ Unverified"
            chunk_ref = g["chunk_id"] or "—"
            line = f"{i}. {g['source']}  [{status}]  (ref: {chunk_ref})"
            p = doc.add_paragraph(line)
            _set_font(p, font_name=fmt.font_name, size=fmt.font_size_pt - 2)

    abs_path = os.path.abspath(output_path)
    doc.save(abs_path)

    if lawyer_notes_text.strip():
        stem = Path(output_path).stem
        parent = Path(output_path).parent
        notes_path = str(parent / f"{stem}_notes.docx")
        _write_lawyer_notes_docx(lawyer_notes_text, notes_path)

    return abs_path


# ---------------------------------------------------------------------------
# Affidavit sub-document writer
# ---------------------------------------------------------------------------

def write_affidavit_docx(state: SeniorCounselState, output_path: str) -> str:
    """
    Write state["affidavit_output"] to a standalone affidavit_evidence.docx.

    The affidavit is the complainant's examination-in-chief in S.138 summary
    trials (S.143 NI Act). It is a separate document from the complaint —
    never a section within it.

    Returns the absolute path of the written file.
    """
    affidavit_text = state.get("affidavit_output") or ""
    if not affidavit_text.strip():
        return ""

    parties = state.get("parties") or {}
    complainant_name = (
        parties.get("complainant")
        or parties.get("plaintiff")
        or parties.get("petitioner")
        or "Complainant"
    )

    soul = state.get("lawyer_soul") or {}
    soul_dict = soul if isinstance(soul, dict) else {}
    jurisdiction = state.get("jurisdiction") or ""
    matter_type = state.get("matter_type") or ""

    spec = load_spec(jurisdiction, matter_type, soul=soul_dict)
    fmt = spec.formatting

    doc = Document()
    _apply_page_setup(doc, spec)

    court_header = soul_dict.get("court_name_formal") or _format_court_header(jurisdiction)
    if court_header:
        _add_court_header(doc, court_header, font_name=fmt.font_name, font_size_pt=fmt.font_size_pt)
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"EVIDENCE BY AFFIDAVIT OF {complainant_name.upper()} (CW-1)")
    run.bold = True
    run.font.name = fmt.font_name
    run.font.size = Pt(fmt.font_size_pt)
    doc.add_paragraph()

    for para_text in [p.strip() for p in re.split(r"\n\s*\n", affidavit_text) if p.strip()]:
        p = doc.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_font(p, font_name=fmt.font_name, size=fmt.font_size_pt)
        _set_paragraph_spacing(p, line_spacing=fmt.line_spacing)

    abs_path = os.path.abspath(output_path)
    doc.save(abs_path)
    return abs_path


# -----------------------------------------------------------------------
# Formatting helpers
# -----------------------------------------------------------------------


def _set_font(
    para,
    font_name: str = "Bookman Old Style",
    size: int = 14,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Apply font settings to every run in a paragraph."""
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic


def _set_paragraph_spacing(para, line_spacing: str = "single") -> None:
    """Apply line spacing from the court spec's line_spacing rule."""
    fmt = para.paragraph_format
    fmt.space_after = Pt(0)
    if line_spacing == "double":
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif line_spacing == "1.5x":
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:  # "single" or any unrecognised value
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _apply_page_setup(doc: Document, spec: CourtDraftSpec) -> None:
    """Set page size and margins from the court spec."""
    fmt = spec.formatting
    for section in doc.sections:
        if fmt.paper == "Legal":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14.0)
        else:  # A4
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        section.left_margin = Inches(fmt.left_margin_inches)
        section.right_margin = Inches(fmt.right_margin_inches)
        section.top_margin = Inches(fmt.top_margin_inches)
        section.bottom_margin = Inches(fmt.bottom_margin_inches)
