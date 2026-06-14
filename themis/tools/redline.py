"""
Word-compatible tracked-changes redlining for Themis.

Diffs original .docx paragraph-by-paragraph against revised_text,
injects OOXML <w:del>/<w:ins> so Word shows tracked changes natively.

WHY lxml direct injection: python-docx has no tracked-changes API;
we reach into paragraph._p and build the OOXML namespace-aware elements
directly — the same encoding Word uses internally.
"""
from __future__ import annotations
import difflib, uuid
from datetime import datetime, timezone
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_AUTHOR = "Themis"
_DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _rev_id() -> str:
    return str(uuid.uuid4().int)[:8]


def _make_del_run(text: str) -> OxmlElement:
    el = OxmlElement("w:del")
    el.set(qn("w:id"), _rev_id())
    el.set(qn("w:author"), _AUTHOR)
    el.set(qn("w:date"), _DATE)
    r = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    r.append(dt)
    el.append(r)
    return el


def _make_ins_run(text: str) -> OxmlElement:
    el = OxmlElement("w:ins")
    el.set(qn("w:id"), _rev_id())
    el.set(qn("w:author"), _AUTHOR)
    el.set(qn("w:date"), _DATE)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    el.append(r)
    return el


def _patch_paragraph(para, old_text: str, new_text: str) -> None:
    p = para._p
    for child in list(p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "del", "ins"):
            p.remove(child)
    old_words = old_text.split()
    new_words = new_text.split()
    for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(
        None, old_words, new_words, autojunk=False
    ).get_opcodes():
        if tag == "equal":
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = " ".join(old_words[i1:i2]) + " "
            r.append(t)
            p.append(r)
        elif tag == "replace":
            p.append(_make_del_run(" ".join(old_words[i1:i2])))
            p.append(_make_ins_run(" ".join(new_words[j1:j2])))
        elif tag == "delete":
            p.append(_make_del_run(" ".join(old_words[i1:i2])))
        elif tag == "insert":
            p.append(_make_ins_run(" ".join(new_words[j1:j2])))


def write_redline_docx(original_path: str, revised_text: str, output_path: str) -> str:
    """
    Produce a Word-compatible redlined .docx.

    Args:
        original_path: Path to original .docx.
        revised_text:  Revised draft as plain string (paragraphs separated by newlines).
        output_path:   Destination path.

    Returns:
        Absolute path of the written file.
    """
    doc = Document(original_path)
    revised_paras = [p for p in revised_text.split("\n") if p.strip()]
    for i, para in enumerate(doc.paragraphs):
        old_text = para.text.strip()
        if not old_text:
            continue
        new_text = revised_paras[i].strip() if i < len(revised_paras) else ""
        if old_text != new_text:
            _patch_paragraph(para, old_text, new_text)
    for j in range(len(doc.paragraphs), len(revised_paras)):
        new_para = doc.add_paragraph()
        _patch_paragraph(new_para, "", revised_paras[j].strip())
    doc.save(output_path)
    return str(Path(output_path).resolve())
