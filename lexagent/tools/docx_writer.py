# Court-ready .docx formatter using python-docx.
#
# Output contract (court-fileable):
#   - Formal court header: all-caps, bold, centered plain paragraph
#   - Cause title: party labels driven by matter_type (Complainant/Accused, not Petitioner/Respondent)
#   - Body: draft paragraphs (double-spaced, justified) — lawyer notes stripped
#   - Citations appendix: grounded citations with source chunk reference
#   - lawyer_notes.docx: Plain English Summary + Risk Assessment routed here, NOT in the filing
#
# Formatting: Times New Roman 12pt, 1.5" left margin — standard Indian district court filing.

from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document  # type: ignore[import]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
from docx.shared import Inches, Pt  # type: ignore[import]

from lexagent.state import LexState

# ---------------------------------------------------------------------------
# Party label map — keyed by matter_type value from LexState.
# WHY: "Petitioner v. Respondent" is writ-petition vocabulary.
# An S.138 criminal complaint must say "Complainant ... Accused" or the
# court officer will flag a format defect at the filing counter.
# Default falls back to Petitioner/Respondent for unknown matter types.
# ---------------------------------------------------------------------------
PARTY_LABELS: dict[str, tuple[str, str]] = {
    "s138_complaint":         ("Complainant", "Accused"),
    "criminal_complaint":     ("Complainant", "Accused"),
    "bail_application":       ("Accused", "State"),
    "writ_petition":          ("Petitioner", "Respondent"),
    "plaint":                 ("Plaintiff", "Defendant"),
    "written_statement":      ("Plaintiff", "Defendant"),
    "injunction_application": ("Plaintiff/Applicant", "Defendant/Respondent"),
    "legal_notice":           ("Sender", "Addressee"),
    "civil_suit":             ("Plaintiff", "Defendant"),
}
_DEFAULT_PARTY_LABELS = ("Petitioner", "Respondent")


def _party_labels(matter_type: str | None) -> tuple[str, str]:
    if not matter_type:
        return _DEFAULT_PARTY_LABELS
    key = (matter_type or "").lower().strip().replace(" ", "_")
    # Exact match first, then substring scan
    if key in PARTY_LABELS:
        return PARTY_LABELS[key]
    for k, v in PARTY_LABELS.items():
        if k in key or key in k:
            return v
    return _DEFAULT_PARTY_LABELS


def _filer_party(matter_type: str | None, parties: dict) -> str:
    """Return the filer-side party name from the parties dict using the correct key."""
    label_filer, _ = _party_labels(matter_type)
    label_key = label_filer.lower()
    # Try exact key match, then alias keys
    return (
        parties.get(label_key)
        or parties.get("plaintiff")
        or parties.get("complainant")
        or parties.get("petitioner")
        or parties.get("accused")
        or label_filer
    )


def _opposing_party(matter_type: str | None, parties: dict) -> str:
    """Return the opposing-side party name from the parties dict."""
    _, label_opposing = _party_labels(matter_type)
    label_key = label_opposing.lower()
    return (
        parties.get(label_key)
        or parties.get("defendant")
        or parties.get("accused")
        or parties.get("respondent")
        or label_opposing
    )


# ---------------------------------------------------------------------------
# Court header renderer
# ---------------------------------------------------------------------------

def _format_court_header(jurisdiction: str) -> str:
    """
    Convert a jurisdiction string to a formal court header.

    Examples:
      "GBN CJM" → "IN THE COURT OF THE CHIEF JUDICIAL MAGISTRATE, GAUTAM BUDH NAGAR"
      "Delhi High Court" → "IN THE HIGH COURT OF DELHI AT NEW DELHI"
      "IN THE COURT OF ..." → returned unchanged (already formal)

    If the string already starts with "IN THE", it is returned as-is, uppercased.
    Otherwise a best-effort expansion is applied; raw string is uppercased as fallback.
    """
    j = jurisdiction.strip()
    if j.upper().startswith("IN THE"):
        return j.upper()

    # Common shorthand expansions
    expansions = {
        "gbncjm":             "IN THE COURT OF THE CHIEF JUDICIAL MAGISTRATE, GAUTAM BUDH NAGAR AT GREATER NOIDA",
        "gbn cjm":            "IN THE COURT OF THE CHIEF JUDICIAL MAGISTRATE, GAUTAM BUDH NAGAR AT GREATER NOIDA",
        "delhi hc":           "IN THE HIGH COURT OF DELHI AT NEW DELHI",
        "delhi high court":   "IN THE HIGH COURT OF DELHI AT NEW DELHI",
        "bombay hc":          "IN THE HIGH COURT OF JUDICATURE AT BOMBAY",
        "bombay high court":  "IN THE HIGH COURT OF JUDICATURE AT BOMBAY",
        "madras hc":          "IN THE HIGH COURT OF JUDICATURE AT MADRAS",
        "allahabad hc":       "IN THE HIGH COURT OF JUDICATURE AT ALLAHABAD",
        "supreme court":      "IN THE SUPREME COURT OF INDIA",
    }
    key = j.lower().strip()
    if key in expansions:
        return expansions[key]

    # If it looks like "Chief Judicial Magistrate, XYZ" wrap it
    if "chief judicial magistrate" in key or "cjm" in key:
        return f"IN THE COURT OF THE {j.upper()}"
    if "high court" in key:
        return f"IN THE {j.upper()}"
    if "sessions court" in key or "sessions judge" in key:
        return f"IN THE COURT OF THE {j.upper()}"

    # Fallback: uppercase the raw string — still looks formal
    return j.upper()


def _add_court_header(doc: Document, court_header: str) -> None:
    """Write the court header as a bold, centered plain paragraph (not a heading style)."""
    for line in court_header.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Draft splitter — separates court-filing body from lawyer working notes
# ---------------------------------------------------------------------------

def _split_draft(draft: str) -> tuple[str, str]:
    """
    Split the LLM draft at the first '---' separator line.

    Returns (filing_body, lawyer_notes).
    - filing_body:    everything before the separator → goes into the .docx filing
    - lawyer_notes:   everything after the separator  → goes into lawyer_notes.docx

    The separator pattern is a line that is exactly '---' (possibly with surrounding
    whitespace), which is the convention used in base_system.md to delimit the
    Plain English Summary and Risk Assessment from the legal document body.
    """
    # Match a line that is ONLY dashes (3+), possibly with surrounding whitespace
    parts = re.split(r"\n[ \t]*---+[ \t]*\n", draft, maxsplit=1)
    if len(parts) == 2:
        return parts[0].rstrip(), parts[1].strip()
    return draft, ""


# ---------------------------------------------------------------------------
# Lawyer notes writer — separate .docx for internal use only
# ---------------------------------------------------------------------------

def _write_lawyer_notes_docx(notes_text: str, notes_path: str) -> str:
    """Write lawyer notes (Plain English Summary + Risk flags) to a separate .docx."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("LAWYER'S WORKING NOTES — NOT FOR FILING")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    for para_text in [p.strip() for p in re.split(r"\n\s*\n", notes_text) if p.strip()]:
        p = doc.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_font(p)
        _set_paragraph_spacing(p)

    abs_path = os.path.abspath(notes_path)
    doc.save(abs_path)
    return abs_path


# ---------------------------------------------------------------------------
# Main public function
# ---------------------------------------------------------------------------

def write_docx(state: LexState, output_path: str) -> str:
    """
    Write state["draft_output"] to a court-ready .docx file.

    - Splits draft at '---': filing body goes into output_path; lawyer notes
      go into a sibling file at <stem>_notes.docx.
    - Uses correct party labels for the matter type.
    - Renders a formal court header from the jurisdiction string.
    - Removes the "Generated by LexAgent" footer from the filing .docx.

    Returns the resolved absolute path of the filing .docx.
    """
    doc = Document()

    # ----------------------------------------------------------------
    # Page margins — standard Indian court filing: 1.5" left, 1" others
    # ----------------------------------------------------------------
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    matter_type = state.get("matter_type") or ""
    parties = state.get("parties") or {}
    jurisdiction = state.get("jurisdiction") or ""

    # ----------------------------------------------------------------
    # Court header — formal all-caps centered paragraph, not a heading
    # WHY: Try SOUL.md court_name_formal first; fall back to jurisdiction string.
    # This matches the human filing's "IN THE COURT OF THE CJM, GBN AT GREATER NOIDA"
    # first line exactly.
    # ----------------------------------------------------------------
    court_header = ""
    soul = state.get("lawyer_soul") or {}
    if isinstance(soul, dict):
        court_header = soul.get("court_name_formal") or ""
    if not court_header and jurisdiction:
        court_header = _format_court_header(jurisdiction)
    if court_header:
        _add_court_header(doc, court_header)
        doc.add_paragraph()  # spacer after header

    # ----------------------------------------------------------------
    # Blank case number line — court fills this at filing counter
    # ----------------------------------------------------------------
    case_no_p = doc.add_paragraph()
    case_no_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = case_no_p.add_run("COMPLAINT CASE NO. _____ OF ____")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    doc.add_paragraph()  # spacer

    # ----------------------------------------------------------------
    # Cause title — uses matter-type-appropriate party labels
    # WHY: Petitioner/Respondent is writ vocabulary. An S.138 criminal
    # complaint must read Complainant / Accused.
    # ----------------------------------------------------------------
    filer_label, opposing_label = _party_labels(matter_type)
    filer_name = _filer_party(matter_type, parties)
    opposing_name = _opposing_party(matter_type, parties)

    filer_p = doc.add_paragraph()
    filer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = filer_p.add_run(f"{filer_name}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    filer_label_p = doc.add_paragraph()
    filer_label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = filer_label_p.add_run(f"... {filer_label.upper()}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    versus_p = doc.add_paragraph()
    versus_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = versus_p.add_run("VERSUS")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    opp_p = doc.add_paragraph()
    opp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = opp_p.add_run(f"{opposing_name}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    opp_label_p = doc.add_paragraph()
    opp_label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = opp_label_p.add_run(f"... {opposing_label.upper()}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # ----------------------------------------------------------------
    # Split draft at '---' to separate filing body from lawyer notes
    # WHY: base_system.md instructs the LLM to append Plain English Summary
    # and Risk Assessment after a '---' separator. These are internal tools
    # for the lawyer — they must never appear in a court filing.
    # ----------------------------------------------------------------
    raw_draft = state.get("draft_output") or ""
    filing_body, lawyer_notes_text = _split_draft(raw_draft)

    # ----------------------------------------------------------------
    # Draft body — split on double newlines into paragraphs
    # ----------------------------------------------------------------
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", filing_body) if p.strip()]
    grounded = state.get("grounded_citations") or []

    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_font(p)
        _set_paragraph_spacing(p)

    # ----------------------------------------------------------------
    # Citations appendix — grounded citations only (no lawyer flags)
    # ----------------------------------------------------------------
    if grounded:
        doc.add_page_break()
        heading_p = doc.add_paragraph()
        heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading_p.add_run("LIST OF CITATIONS")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        for i, g in enumerate(grounded, start=1):
            status = "✓ Verified" if g["verified"] else "⚠ Unverified"
            chunk_ref = g["chunk_id"] or "—"
            line = f"{i}. {g['source']}  [{status}]  (ref: {chunk_ref})"
            p = doc.add_paragraph(line)
            _set_font(p, size=10)

    # ----------------------------------------------------------------
    # Save the filing .docx — NO footer, NO matter ID, NO LexAgent branding
    # ----------------------------------------------------------------
    abs_path = os.path.abspath(output_path)
    doc.save(abs_path)

    # ----------------------------------------------------------------
    # Write lawyer notes to a sibling file if there is any content
    # ----------------------------------------------------------------
    if lawyer_notes_text.strip():
        stem = Path(output_path).stem
        parent = Path(output_path).parent
        notes_path = str(parent / f"{stem}_notes.docx")
        _write_lawyer_notes_docx(lawyer_notes_text, notes_path)

    return abs_path


# -----------------------------------------------------------------------
# Formatting helpers
# -----------------------------------------------------------------------


def _set_font(
    para,
    size: int = 12,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Apply font settings to every run in a paragraph."""
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic


def _set_paragraph_spacing(para) -> None:
    """Double-space body paragraphs with 12pt spacing after."""
    fmt = para.paragraph_format
    fmt.space_after = Pt(12)
    fmt.line_spacing = Pt(24)  # ~double spacing for 12pt font
